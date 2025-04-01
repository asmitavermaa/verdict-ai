from flask import Flask, request, jsonify, render_template, session, send_file

import PyPDF2
import os 
from dotenv import load_dotenv
from io import BytesIO
import tempfile
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uuid
import datetime
from multiagent import questioner, tone_analyzer, summarizer


load_dotenv()

# Initialize Flask app
app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', 'default-secret-key')

general_context = ""
doc_chat_context = ""

# Store document text globally (in a real app, you'd use a database or session)
document_cache = {}
pdf_cache = {}  # Store PDF files for viewing
draft_cache = {}  # Store generated drafts

CATEGORY_METRICS = {
    'Legal Notice': [
        'Severity Score', 'Violations & Broken Rules', 'Legal Consequences', 'Actionable Steps',
        'Urgency Detection', 'Tone Analysis', 'Recommended Actions'
    ],
    'Ownership Documents': [
        'Ownership Rights & Obligations', 'Transfer, Leasing, Sale, Mortgaging Clauses',
        'Financial Liabilities', 'Terms & Conditions', 'Important Dates', 'Document Validity', 'Summary Type'
    ],
    'Contracts & Agreements': [
        'Parties Involved & Roles', 'Terms & Conditions', 'Termination Clauses', 'Penalties for Breach',
        'Severity Score', 'Obligations & Rights', 'Actionable Steps'
    ],
    'Financial Documents': [
        'Financial Obligations', 'Coverage Details', 'Deadlines & Payment Schedules',
        'Legal Implications', 'Severity Score', 'Urgency Detection', 'Risk Analysis'
    ],
    'Terms & Conditions / Privacy Policies': [
        'User Rights & Restrictions', 'Data Usage & Privacy Clauses', 'Liability Clauses',
        'Termination & Suspension Rules', 'Severity Score', 'Personal Implications', 'Suggested Actions'
    ],
    'Intellectual Property Documents': [
        'Ownership & Usage Rights', 'Infringement Clauses', 'Exclusivity & Licensing Terms',
        'Penalties for Violation', 'Severity Score', 'Urgency Detection', 'Recommended Actions'
    ],
    'Criminal Offense Documents': [
        'Charges Filed', 'Potential Penalties', 'Required Actions', 'Severity Score',
        'Urgency Detection', 'Tone Analysis', 'Suggested Actions'
    ],
    'Regulatory Compliance Documents': [
        'Compliance Requirements', 'Penalties for Non-Compliance', 'Renewal Deadlines & Conditions',
        'Guidelines for Rectification', 'Severity Score', 'Urgency Detection', 'Recommended Actions'
    ],
    'Employment Documents': [
        'Terms of Employment', 'Termination Conditions', 'Confidentiality Clauses',
        'Breach Consequences', 'Severity Score', 'Urgency Detection', 'Suggested Actions'
    ],
    'Court Judgments & Legal Precedents': [
        'Summary of Judgment', 'Legal Basis', 'Potential Consequences',
        'Severity Score', 'Urgency Detection', 'Recommended Actions'
    ]
}

# Document templates for different draft types
DRAFT_TEMPLATES = {
    'Legal Notice Response': {
        'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.25, 'right': 1.25},
        'header_format': {'font': 'Times New Roman', 'size': 12, 'bold': True, 'align': 'center'},
        'body_format': {'font': 'Times New Roman', 'size': 12, 'align': 'left'},
        'signature_format': {'font': 'Times New Roman', 'size': 12, 'align': 'left'},
        'date_format': '%B %d, %Y',  # Example: January 1, 2023
        'includes_header': True,
        'includes_date': True,
        'includes_signature': True
    },
    'Contract Response': {
        'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.25, 'right': 1.25},
        'header_format': {'font': 'Arial', 'size': 12, 'bold': True, 'align': 'center'},
        'body_format': {'font': 'Arial', 'size': 11, 'align': 'left'},
        'signature_format': {'font': 'Arial', 'size': 11, 'align': 'left'},
        'date_format': '%d/%m/%Y',  # Example: 01/01/2023
        'includes_header': True,
        'includes_date': True,
        'includes_signature': True
    },
    'General Letter': {
        'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.25, 'right': 1.25},
        'header_format': {'font': 'Calibri', 'size': 12, 'bold': True, 'align': 'left'},
        'body_format': {'font': 'Calibri', 'size': 11, 'align': 'left'},
        'signature_format': {'font': 'Calibri', 'size': 11, 'align': 'left'},
        'date_format': '%B %d, %Y',  # Example: January 1, 2023
        'includes_header': True,
        'includes_date': True,
        'includes_signature': True
    },
    'Legal Memo': {
        'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.25, 'right': 1.25},
        'header_format': {'font': 'Times New Roman', 'size': 14, 'bold': True, 'align': 'center'},
        'body_format': {'font': 'Times New Roman', 'size': 12, 'align': 'left'},
        'signature_format': {'font': 'Times New Roman', 'size': 12, 'align': 'left'},
        'date_format': '%B %d, %Y',  # Example: January 1, 2023
        'includes_header': True,
        'includes_date': True,
        'includes_signature': False
    }
}

# Map document categories to draft templates
CATEGORY_TO_TEMPLATE = {
    'Legal Notice': 'Legal Notice Response',
    'Contracts & Agreements': 'Contract Response',
    'Ownership Documents': 'Legal Memo',
    'Financial Documents': 'Legal Memo',
    'Terms & Conditions / Privacy Policies': 'Legal Memo',
    'Intellectual Property Documents': 'Legal Memo',
    'Criminal Offense Documents': 'Legal Notice Response',
    'Regulatory Compliance Documents': 'Legal Memo',
    'Employment Documents': 'Legal Memo',
    'Court Judgments & Legal Precedents': 'Legal Memo',
    'default': 'General Letter'
}

@app.route('/')
def index():
    # Generate a unique session ID if not exists
    if 'session_id' not in session:
        session['session_id'] = os.urandom(16).hex()
    return render_template('index.html')

@app.route('/general_chat.html')
def general_chat():
    return render_template('general_chat.html')

def extract_text_from_pdf(pdf_file):
    try:
        # Create a BytesIO object to avoid file seeking issues
        pdf_content = BytesIO(pdf_file.read())
        # Reset the file pointer for future operations
        pdf_file.seek(0)
        
        reader = PyPDF2.PdfReader(pdf_content)
        text = ''
        for page_num in range(len(reader.pages)):
            text += reader.pages[page_num].extract_text()
        return text
    except Exception as e:
        app.logger.error(f"PDF extraction error: {str(e)}")
        return str(e)


@app.route('/classify', methods=['POST'])
def classify_document():
    if 'document' not in request.files:
        return jsonify({'error': 'No PDF file uploaded'}), 400

    pdf_file = request.files['document']

    # Save the PDF file in memory for later viewing
    session_id = session.get('session_id', os.urandom(16).hex())
    pdf_content = pdf_file.read()
    pdf_cache[session_id] = pdf_content

    # Reset file pointer
    pdf_file.seek(0)

    document_text = extract_text_from_pdf(pdf_file)

    if not document_text:
        return jsonify({'error': 'Failed to extract text from PDF'}), 400

    document_cache[session_id] = document_text

    try:
        # Use WorqHat-based classifier
        from worqhat_utils import DocumentClassifier
        document_classifier = DocumentClassifier()
        category = document_classifier.classify(document_text)
        return jsonify({'category': category})
    except Exception as e:
        app.logger.error(f"Classification error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/process', methods=['POST'])
def process_document():
    if 'document' not in request.files or 'category' not in request.form:
        return jsonify({'error': 'Document file or category is missing'}), 400

    pdf_file = request.files['document']
    category = request.form['category']
    document_text = extract_text_from_pdf(pdf_file)

    if not document_text or not category:
        return jsonify({'error': 'Document text or category is missing'}), 400

    session_id = session.get('session_id', os.urandom(16).hex())
    document_cache[session_id] = document_text

    try:
        # Use WorqHat-based document processor
        from worqhat_utils import DocumentProcessor
        document_processor = DocumentProcessor()
        summary = document_processor.get_summary(document_text)
        key_phrases = document_processor.extract_key_phrases(document_text)

        return jsonify({
            'summary': summary,
            'key_phrases': key_phrases,
            'document_text': document_text[:200] + '...' if len(document_text) > 200 else document_text
        })
    except Exception as e:
        app.logger.error(f"Processing error: {str(e)}")
        return jsonify({'error': str(e)}), 500




@app.route('/chat', methods=['POST'])
def chat():
    global doc_chat_context
    data = request.json
    user_message = data.get('message')
    category = data.get('category')
    detailed_analysis = data.get('detailed_analysis', False)
    generate_draft = data.get('generate_draft', False)
    draft_instructions = data.get('draft_instructions', '')

    if not user_message:
        return jsonify({'error': 'Message is required'}), 400

    session_id = session.get('session_id')
    document_text = document_cache.get(session_id, '')

    if not document_text:
        return jsonify({'error': 'No document found. Please process a document first.'}), 400

    try:
        # Use up to 3000 characters of document context
        document_context = document_text[:3000]

        # Build the prompt
        system_prompt = f"""You are a legal assistant specializing in {category} documents.
You will answer only questions related to the document and not external questions.
Document text (truncated if needed):
{document_context}

Previous chat:
{doc_chat_context}
"""

        if detailed_analysis:
            system_prompt += "Provide a detailed analysis with legal references and thorough explanations.\n"
        else:
            system_prompt += "Provide concise, clear answers focused on key legal points.\n"

        if generate_draft:
            draft_id = generate_document_draft(user_message, draft_instructions, category, document_context)
            return jsonify({
                'response': "I've prepared a draft document based on your instructions. You can download it using the link below.",
                'draft_id': draft_id
            })

        system_prompt += "If you cannot find the answer in the document, clearly say so. Use **bold** for important points."

        # Use WorqHat API for better document analysis
        from worqhat_utils import respond_to_query
        bot_response = respond_to_query(user_message, context=system_prompt)

        doc_chat_context += f"\nUser: {user_message}\nBot: {bot_response}\n"

        return jsonify({'response': bot_response})

    except Exception as e:
        app.logger.error(f"Chat error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/general_chat', methods=['POST'])
def general_chat_api():
    global general_context
    data = request.json
    user_message = data.get('message')
    detailed_analysis = data.get('detailed_analysis', False)

    if not user_message:
        return jsonify({'error': 'Message is required'}), 400

    try:
        if detailed_analysis:
            # Use WorqHat API for detailed analysis
            response = respond_to_query(f"Provide a detailed legal analysis with reasoning for this question: {user_message}", context=general_context)
            # For reasoning, we'll extract key points from the response
            reasoning_prompt = f"Extract 3-5 key legal reasoning points from this response as a list: {response}"
            reasoning = respond_to_query(reasoning_prompt).split('\n')
            general_context += f"\nUser: {user_message}\nSenior Lawyer: {response}\n"
            return jsonify({'response': response, 'reasoning': reasoning})
        else:
            system_prompt = f"""
You are a knowledgeable legal assistant who can provide general information about legal topics.
You are not a lawyer and should clarify that your responses do not constitute legal advice.
Recommend consulting a qualified attorney for specific legal situations.
Use **bold** for important points and be clear and organized.

Context:
{general_context}
"""

            # Use WorqHat API for general legal chat
            from worqhat_utils import respond_to_query
            bot_response = respond_to_query(user_message, context=system_prompt)
            general_context += f"\nUser: {user_message}\nBot: {bot_response}\n"
            return jsonify({'response': bot_response, 'reasoning': []})

    except Exception as e:
        app.logger.error(f"General chat error: {str(e)}")
        return jsonify({'error': str(e)}), 500

def generate_document_draft(message, instructions, category, document_context):
    template_name = CATEGORY_TO_TEMPLATE.get(category, CATEGORY_TO_TEMPLATE['default'])
    template = DRAFT_TEMPLATES[template_name]

    prompt = f"""You are a professional legal document drafter. Create a formal response document based on the following instructions:

Instructions: {instructions}

This is related to a {category} document. Here's the relevant context from the document:
{document_context[:1500]}

Your draft should be well-structured and professionally formatted. Include:
1. A clear header/title
2. Today's date ({datetime.datetime.now().strftime(template['date_format'])})
3. Appropriate salutation
4. Well-organized body content
5. Proper closing
6. Signature line

Do not include explanatory notes. Just return the document content.
"""

    try:
        # Use WorqHat API for document draft generation
        from worqhat_utils import respond_to_query
        draft_content = respond_to_query(message, context=prompt)

        doc = create_formatted_document(draft_content, template)
        draft_id = str(uuid.uuid4())
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)

        draft_cache[draft_id] = {
            'path': temp_file.name,
            'filename': f"Legal_Draft_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
        }

        return draft_id
    except Exception as e:
        app.logger.error(f"Draft generation error: {str(e)}")
        raise


def generate_general_draft(message, instructions):
    """Generate a formatted document draft for general legal inquiries"""
    
    # Use the general letter template for general drafts
    template = DRAFT_TEMPLATES['General Letter']
    
    # Create a prompt for the draft generation
    prompt = f"""You are a professional legal document drafter. Create a formal document based on the following instructions:

Instructions: {instructions}

Your draft should be well-structured and professionally formatted. Include:
1. A clear header/title
2. Today's date ({datetime.datetime.now().strftime(template['date_format'])})
3. Appropriate salutation (if applicable)
4. Well-organized body content
5. Proper closing
6. Signature line (if applicable)

Format the content as if it were going to be printed on letterhead. Do not include any explanatory text or notes - just the actual document content.
"""

    try:
        # Generate the draft content using WorqHat API directly
        from worqhat_utils import worqhat_client
        draft_content = worqhat_client.generate_text(message, context=prompt)

        # Create a formatted Word document
        doc = create_formatted_document(draft_content, template)

        # Save the document to a temporary file
        draft_id = str(uuid.uuid4())
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)

        # Store the file path in the draft cache
        draft_cache[draft_id] = {
            'path': temp_file.name,
            'filename': f"Legal_Draft_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
        }

        return draft_id

    except Exception as e:
        app.logger.error(f"Draft generation error: {str(e)}")
        raise


def create_formatted_document(content, template):
    """Create a properly formatted Word document based on the template and content"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(template['margins']['top'])
        section.bottom_margin = Inches(template['margins']['bottom'])
        section.left_margin = Inches(template['margins']['left'])
        section.right_margin = Inches(template['margins']['right'])
    
    # Split content into lines
    lines = content.split('\n')
    
    # Process each line and apply appropriate formatting
    current_section = 'header'
    
    for line in lines:
        line = line.strip()
        if not line:
            # Add empty paragraph for spacing
            doc.add_paragraph()
            continue
        
        # Determine the section based on content
        if current_section == 'header' and template['includes_date'] and any(month in line.lower() for month in ['january', 'february', 'march', 'april', 'may', 'june', 'july', 'august', 'september', 'october', 'november', 'december']) and any(str(day) in line for day in range(1, 32)):
            current_section = 'date'
        elif current_section in ['header', 'date'] and any(salutation in line.lower() for salutation in ['dear', 'to whom', 'attention', 'attn', 're:', 'subject:']):
            current_section = 'salutation'
        elif current_section in ['header', 'date', 'salutation'] and any(closing in line.lower() for closing in ['sincerely', 'regards', 'truly', 'thank you', 'best', 'respectfully']):
            current_section = 'signature'
        elif current_section in ['header', 'date', 'salutation'] and len(line) > 20:
            current_section = 'body'
        
        # Apply formatting based on the section
        p = doc.add_paragraph()
        
        if current_section == 'header':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if template['header_format']['align'] == 'center' else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.name = template['header_format']['font']
            run.font.size = Pt(template['header_format']['size'])
            run.font.bold = template['header_format']['bold']
        
        elif current_section == 'date':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.name = template['body_format']['font']
            run.font.size = Pt(template['body_format']['size'])
        
        elif current_section == 'body':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.name = template['body_format']['font']
            run.font.size = Pt(template['body_format']['size'])
        
        elif current_section == 'signature':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.name = template['signature_format']['font']
            run.font.size = Pt(template['signature_format']['size'])
    
    return doc


@app.route('/download-draft/<draft_id>', methods=['GET'])
def download_draft(draft_id):
    """Download a generated draft document"""
    
    if draft_id not in draft_cache:
        return jsonify({'error': 'Draft not found'}), 404
    
    draft_info = draft_cache[draft_id]
    
    try:
        return send_file(
            draft_info['path'],
            as_attachment=True,
            download_name=draft_info['filename'],
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        app.logger.error(f"Draft download error: {str(e)}")
        return jsonify({'error': 'Error downloading draft'}), 500


@app.route('/view-document', methods=['GET'])
def view_document():
    session_id = session.get('session_id')
    if not session_id or session_id not in pdf_cache:
        return jsonify({'error': 'No document found'}), 404
    
    # Create a temporary file to serve
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    temp_file.write(pdf_cache[session_id])
    temp_file.close()
    
    return send_file(temp_file.name, mimetype='application/pdf', as_attachment=False)


if __name__ == '__main__':
    app.run(debug=True)

